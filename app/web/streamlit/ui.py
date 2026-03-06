"""
JusReport — aplicação unificada (sem FastAPI).
Todo o pipeline roda direto no Streamlit:
  PDF → extração de texto → Gemini (9 agentes) → DOCX → download
"""

import os
import sys
import io
import time
import traceback
import base64
from datetime import date
from io import BytesIO
from typing import Optional

# ── path ──
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

import pandas as pd
import streamlit as st
from dotenv import load_dotenv

# ── dependências do pipeline ──
import pdfplumber
import google.generativeai as genai
from docx import Document
from docx.shared import Pt

# ── banco de dados ──
from app.utils.db import (  # type: ignore
    salvar_processo, listar_processos, registrar_relatorio,
    excluir_processo, REL_DIR,
)

# ═══════════════════════════════════════════════
# CONFIGURAÇÃO
# ═══════════════════════════════════════════════

load_dotenv(os.path.join(PROJECT_ROOT, ".env"))

RELATORIOS_DIR = REL_DIR
os.makedirs(RELATORIOS_DIR, exist_ok=True)

GEMINI_API_KEY    = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL_TEXT = os.getenv("GEMINI_MODEL_TEXT", "gemini-2.5-pro")
MAX_PDF_CHARS     = int(os.getenv("MAX_PDF_CHARS", "30000"))

SUMARIZACOES = [
    "Execução", "Ação de Cobrança", "Ação Monitória",
    "Embargos à Execução", "Reintegração de Posse",
]

HOTSPOT_KEYWORDS = [
    "planilha", "demonstrativo", "cálculo", "calculo",
    "sisbajud", "bacenjud", "bloqueio", "penhora online", "penhora on-line",
]

# ── inicializa Gemini uma única vez ──
@st.cache_resource
def _get_gemini_model():
    if not GEMINI_API_KEY:
        return None
    genai.configure(api_key=GEMINI_API_KEY)
    try:
        return genai.GenerativeModel(GEMINI_MODEL_TEXT)
    except Exception:
        return genai.GenerativeModel("gemini-1.5-pro")

# ═══════════════════════════════════════════════
# PIPELINE — EXTRAÇÃO DE TEXTO
# ═══════════════════════════════════════════════

def _detect_hotspot_pages(text_by_page: list[str]) -> list[int]:
    return [
        idx + 1 for idx, t in enumerate(text_by_page)
        if any(k in (t or "").lower() for k in HOTSPOT_KEYWORDS)
    ]

def _build_global_sample(full_text: str, max_chars: int) -> str:
    total = len(full_text)
    if total <= max_chars:
        return full_text
    part = max(max_chars // 4, 1)
    mid  = total // 2
    return (
        full_text[:part]
        + "\n\n=== TRECHO CENTRAL ===\n\n"
        + full_text[max(0, mid - part // 2): mid + part // 2]
        + "\n\n=== TRECHO FINAL ===\n\n"
        + full_text[-part:]
    )

def extrair_texto_pdf(path: str) -> str:
    text_by_page: list[str] = []
    pages_obj = []

    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages_obj.append(page)
                text_by_page.append(page.extract_text() or "")
    except Exception as e:
        st.warning(f"Erro ao ler PDF: {e}")
        return ""

    full_text = "\n\n".join(text_by_page)
    total     = len(full_text)
    HARD_CAP  = 80_000
    max_chars = min(MAX_PDF_CHARS, HARD_CAP)

    if total <= max_chars:
        return full_text

    # Hotspots
    hotspot_parts: list[str] = []
    for idx, page_text in enumerate(text_by_page):
        tl = (page_text or "").lower()
        if not any(k in tl for k in HOTSPOT_KEYWORDS):
            continue
        pn    = idx + 1
        bloco = [f"\n\n=== PÁGINA RELEVANTE {pn} ===\n\n", page_text]
        try:
            tables = pages_obj[idx].extract_tables() or []
            if tables:
                bloco.append(f"\n\n=== TABELA(S) PÁG. {pn} ===\n\n")
                for t in tables:
                    for row in t:
                        bloco.append(" | ".join(cell or "" for cell in row) + "\n")
        except Exception:
            pass
        hotspot_parts.append("".join(bloco))

    hotspot_text = "".join(hotspot_parts).strip()
    if not hotspot_text:
        return _build_global_sample(full_text, max_chars)

    max_hot  = int(max_chars * 0.6)
    hot_trim = hotspot_text[:max_hot]
    remain   = max_chars - len(hot_trim)
    return hot_trim + "\n\n=== AMOSTRAGEM GLOBAL ===\n\n" + _build_global_sample(full_text, remain)


# ═══════════════════════════════════════════════
# PIPELINE — AGENTES GEMINI
# ═══════════════════════════════════════════════

def _build_tasks(action_type: str) -> list[dict]:
    return [
        {
            "key": "sumario_executivo",
            "title": "1. Sumário Executivo",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Com base EXCLUSIVAMENTE no texto do processo abaixo, produza "# 1. SUMÁRIO EXECUTIVO".
Crie uma tabela: **INDICADOR** | **DADOS EXTRAÍDOS** | **FONTE**

Linhas obrigatórias (use "Não localizado" se ausente):
Processo nº | Vara / Comarca | Data de Distribuição | Classe |
Exequente (Credor) | Executado Principal | Co-executados |
Advogado Credor (atual) | Advogado(s) dos Executados |
Valor da Causa (Inicial) | Valor Atualizado (Exigível) | Título Executivo |
Bem Penhorado / Garantia Principal | Avaliação do Bem | Risco de Prescrição.

Na coluna FONTE indique número do documento, página ou referência.
""",
        },
        {
            "key": "mapa_citacao",
            "title": "2. Mapa de Citação e Partes",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 2. MAPA DE CITAÇÃO E PARTES — AUDITORIA DE NULIDADES".
Para cada parte (executado principal, co-executados, fiadores, terceiros interessados)
crie subseção "## 2.X Nome — Qualificação" com:
**Qualificação** | **Status da Citação** | **Forma** | **Representação** |
**Evento Morte/Sucessão** | **Risco de Nulidade** (BAIXO/MÉDIO/ALTO + justificativa).
Se terceiros ingressaram nos autos, detalhe: qualificação, data, tese, resposta do credor, status.
""",
        },
        {
            "key": "titulo_origem",
            "title": "3. Título e Origem da Dívida",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 3. TÍTULO E ORIGEM DA DÍVIDA".
Tabela: **ASPECTO** | **DESCRIÇÃO** | **FONTE**
Linhas: Título Executivo | Identificação/Nº | Classe Processual | Credor Original |
Devedor Principal | Co-devedores/Avalistas | Data de Distribuição |
Garantia Original | Valor Original | Valor Atualizado | Registro de Imóveis | Observações.
Após a tabela, se houver nota crítica sobre garantia, inclua:
"*NOTA CRÍTICA — [TEMA]: ...*"
""",
        },
        {
            "key": "asset_tracing",
            "title": "4. Rastreamento de Ativos",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 4. RASTREAMENTO DE ATIVOS — ASSET TRACING" com:

## 4.1 Penhoras Efetivadas
Tabela: **BEM** | **DESCRIÇÃO** | **FONTE** — tipo, descrição completa, data do termo,
depositário, status do registro, referência documental.

## 4.2 Avaliação do Bem Penhorado
Tabela: **ITEM** | **DADOS** | **FONTE** — avaliador, data, área, metodologia,
VALOR DA AVALIAÇÃO, homologação, impugnações, relação dívida/garantia, uso do bem.

## 4.3 Diligências de Busca de Bens
Para cada: SISBAJUD/BACENJUD, RENAJUD, INFOJUD, SERASAJUD, outros —
pedido, decisão, resultado, referência. Se ausente: "Não localizado nos trechos analisados."
""",
        },
        {
            "key": "status_expropriatorio",
            "title": "5. Status Expropriatório",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 5. STATUS EXPROPRIATÓRIO — LEILÃO / ADJUDICAÇÃO".
Tabela cronológica: **DATA / ATO** | **DESCRIÇÃO** | **STATUS** | **FONTE**
— pedidos de leilão, decisões, incidentes bloqueadores, designação, arrematação/adjudicação.
Finalize com parágrafo **CONCLUSÃO** em negrito resumindo o estágio atual e o principal obstáculo.
""",
        },
        {
            "key": "defesas_incidentes",
            "title": "6. Defesas e Incidentes",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 6. DEFESAS E INCIDENTES — O PASSIVO".
Para cada incidente: subseção "## 6.X [Nome]" com:
**Data** | **Documento** | **Requerente** | **Tese Principal** |
**Decisão** | **Recurso/Desdobramento** | **Status**.
Inclua: embargos à execução, exceção de pré-executividade, impugnação à avaliação,
substituição de penhora, embargos de terceiro, fraude à execução, desconsideração da PJ,
embargos de declaração relevantes, prescrição intercorrente.
Se não localizar: "Não localizado nos autos digitalizados analisados."
""",
        },
        {
            "key": "red_flags",
            "title": "7. Red Flags e Riscos Críticos",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 7. RED FLAGS E RISCOS CRÍTICOS" com dois grupos:

**GRUPO 1 — RED FLAGS**
**⚠ RED FLAG #N — [TÍTULO] [RISCO ALTO/MÉDIO]**
Descreva: o que foi encontrado, por que representa risco, impacto potencial, referência documental.
Pesquise: disputa de propriedade, penhora não averbada, citação não comprovada,
alteração de matrícula, prescrição intercorrente, autos físicos não digitalizados, outros.

**GRUPO 2 — PONTOS POSITIVOS**
**✔ PONTO POSITIVO #N — [TÍTULO]**
Descreva o fator e seu impacto positivo.
Se nenhum red flag: "Nenhum red flag crítico identificado nos trechos analisados."
""",
        },
        {
            "key": "linha_do_tempo",
            "title": "8. Índice Remissivo — Linha do Tempo",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 8. ÍNDICE REMISSIVO — LINHA DO TEMPO".
Tabela cronológica: **DATA** | **EVENTO** | **DOCUMENTO / FONTE**
— distribuição, citações, migrações, juntadas relevantes, pedidos, decisões,
penhoras, avaliações, hasta pública, substabelecimentos, embargos, último ato.
Datas inexatas: "Mês/Ano (data aproximada)". Prefira listar mais atos que menos.
""",
        },
        {
            "key": "parecer_final",
            "title": "9. Parecer Final e Recomendação",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.
Produza "# 9. PARECER FINAL E RECOMENDAÇÃO DE AQUISIÇÃO".

**SÍNTESE DA ANÁLISE:** parágrafo de 3–5 linhas.

Tabela: **CRITÉRIO** | **AVALIAÇÃO** | **COMENTÁRIO**
Linhas: LIQUIDEZ | CERTEZA | EXIGIBILIDADE | SOLVÊNCIA DO DEVEDOR | GARANTIA

**RECOMENDAÇÃO: [AQUISIÇÃO RECOMENDADA / CONDICIONADA / NÃO RECOMENDADA]**

Se CONDICIONADA:
- CONDIÇÃO 1 — [TÍTULO]: descrição.
- CONDIÇÃO 2 — ...

Finalize com parágrafo em itálico sobre potencial de recuperação e riscos residuais.
Baseie-se APENAS no que consta nos autos.
""",
        },
    ]


def _call_gemini(model, task: dict, base_text: str) -> tuple[str, str]:
    prompt = (
        f"{task['instruction']}\n\n"
        f"=== TEXTO DO PROCESSO ===\n\n\"\"\"{base_text}\"\"\""
    )
    resp = model.generate_content(prompt)
    return task["key"], (resp.text or "").strip()


def gerar_relatorio_md(base_text: str, case_number: str, action_type: str,
                        progress_bar=None) -> str:
    model = _get_gemini_model()
    if not model:
        raise RuntimeError("GEMINI_API_KEY não configurada.")

    tasks    = _build_tasks(action_type)
    sections: dict[str, str] = {}
    n        = len(tasks)

    for i, task in enumerate(tasks):
        try:
            key, text = _call_gemini(model, task, base_text)
            sections[key] = text
        except Exception as e:
            sections[task["key"]] = f"Erro ao processar esta seção: {e}"
        if progress_bar:
            progress_bar.progress(int((i + 1) / n * 100),
                                  text=f"Seção {i+1}/{n} — {task['title']}...")

    meses = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    hoje  = date.today()
    data_str = f"{hoje.day} de {meses[hoje.month-1]} de {hoje.year}"

    header = (
        f"**RELATÓRIO DE PRÉ-AUDITORIA**\n\n"
        f"**LEGAL DUE DILIGENCE — NPL**\n\n"
        f"{action_type} — Processo nº {case_number}\n\n"
        f"*Data de Geração: {data_str} | CONFIDENCIAL*"
    )

    parts = [header]
    for task in tasks:
        txt = sections.get(task["key"], "").strip()
        parts.append(txt if txt
                     else f"# {task['title']}\n\nNão localizado nos trechos analisados.")

    return "\n\n---\n\n".join(parts)


# ═══════════════════════════════════════════════
# PIPELINE — GERAÇÃO DO DOCX
# ═══════════════════════════════════════════════

def gerar_docx(markdown_text: str) -> bytes:
    doc   = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped == "---":
            doc.add_paragraph("─" * 60)
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════
# BANCO DE DADOS — CACHE
# ═══════════════════════════════════════════════

@st.cache_data(ttl=30, show_spinner=False)
def pendentes_df() -> pd.DataFrame:
    rows = listar_processos(status="pendente")
    cols = ["id","nome_cliente","numero_processo","tipo","data_envio","caminho_arquivo"]
    if not rows: return pd.DataFrame(columns=cols)
    df = pd.DataFrame(rows)
    for c in cols:
        if c not in df.columns: df[c] = None
    return df[cols].sort_values("data_envio", ascending=False)

@st.cache_data(ttl=30, show_spinner=False)
def finalizados_df() -> pd.DataFrame:
    rows = listar_processos(status="finalizado")
    cols = ["id","nome_cliente","numero_processo","data_envio","caminho_relatorio"]
    if not rows: return pd.DataFrame(columns=cols)
    df = pd.DataFrame(rows)
    for c in cols:
        if c not in df.columns: df[c] = None
    return df[cols].sort_values("data_envio", ascending=False)

@st.cache_data(ttl=30, show_spinner=False)
def mensal_df() -> pd.DataFrame:
    rows = listar_processos(status=None)
    if not rows: return pd.DataFrame(columns=["Colaborador","Mês/Ano","Processos"])
    df = pd.DataFrame(rows)
    df["data_envio"] = pd.to_datetime(df["data_envio"], errors="coerce")
    df["mes_ano"]    = df["data_envio"].dt.strftime("%m/%Y")
    result = (df.groupby(["nome_cliente","mes_ano"]).size()
                .reset_index(name="quantidade")
                .sort_values("mes_ano", ascending=False))
    result.columns = ["Colaborador","Mês/Ano","Processos"]
    return result

def invalidar():
    pendentes_df.clear(); finalizados_df.clear(); mensal_df.clear()

def excluir_com_arquivo(pid: str, caminho: Optional[str]):
    excluir_processo(pid)
    if caminho and os.path.exists(caminho):
        try: os.remove(caminho)
        except: pass

def _guess_mime(fn: str) -> str:
    return ("application/pdf" if fn.lower().endswith(".pdf") else
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fn.lower().endswith(".docx") else "application/octet-stream")


# ═══════════════════════════════════════════════
# CSS
# ═══════════════════════════════════════════════

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;600;700&family=Outfit:wght@300;400;500;600&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --gold:#C9A84C; --gold-b:#E4C06E; --gold-d:#8A6A1F;
  --navy:#07080F; --p1:#0D0F1A; --p2:#121520; --p3:#181C2A;
  --b1:#1E2236; --b2:#252A40;
  --tx:#E8E4D9; --tx2:#9B97A6; --tx3:#555A70;
  --green:#2E7D5E; --greenl:#3DAA7B;
}

html,body,[data-testid="stAppViewContainer"],.main .block-container{
  background:var(--navy)!important; color:var(--tx)!important;
  font-family:'Outfit',sans-serif!important; max-width:100%!important; padding-top:0!important;
}
[data-testid="stSidebar"]{
  background:var(--p1)!important; border-right:1px solid var(--b1)!important; padding:0!important;
}
[data-testid="stSidebar"]>div:first-child,[data-testid="stSidebarContent"]{padding:0!important;}

[data-testid="stSidebar"] .stRadio>div{gap:2px!important;}
[data-testid="stSidebar"] .stRadio label{
  display:flex!important; align-items:center!important; padding:10px 20px!important;
  margin:0!important; border-radius:0!important; color:var(--tx2)!important;
  font-family:'Outfit',sans-serif!important; font-size:0.82rem!important;
  font-weight:400!important; letter-spacing:0.05em!important; cursor:pointer!important;
  transition:all 0.15s!important; border-left:2px solid transparent!important;
}
[data-testid="stSidebar"] .stRadio label:hover{
  background:rgba(201,168,76,0.06)!important; color:var(--tx)!important;
  border-left-color:var(--gold-d)!important;
}
[data-testid="stSidebar"] [data-testid="stRadioLabel"] p{
  font-family:'Outfit',sans-serif!important; font-size:0.82rem!important;
  color:inherit!important; margin:0!important;
}
[data-testid="stSidebar"] .stRadio input[type="radio"]{display:none!important;}

h1,h2,h3{font-family:'Cormorant Garamond',serif!important; color:var(--tx)!important; letter-spacing:0.02em!important;}

[data-testid="stTextInput"] input{
  background:var(--p3)!important; border:1px solid var(--b2)!important;
  border-radius:6px!important; color:var(--tx)!important;
  font-family:'Outfit',sans-serif!important; font-size:0.88rem!important;
  padding:10px 14px!important; transition:border-color 0.2s!important;
}
[data-testid="stTextInput"] input:focus{
  border-color:var(--gold-d)!important; box-shadow:0 0 0 3px rgba(201,168,76,0.12)!important;
}
[data-testid="stTextInput"] label,[data-testid="stSelectbox"] label,[data-testid="stFileUploader"] label{
  color:var(--tx3)!important; font-size:0.7rem!important; font-weight:500!important;
  letter-spacing:0.1em!important; text-transform:uppercase!important;
}
div[data-baseweb="select"]>div{
  background:var(--p3)!important; border:1px solid var(--b2)!important;
  border-radius:6px!important; color:var(--tx)!important;
}
[data-testid="stFileUploader"] section{
  background:var(--p3)!important; border:1px dashed var(--b2)!important;
  border-radius:8px!important;
}
[data-testid="stFileUploader"] section:hover{border-color:var(--gold-d)!important;}
[data-testid="stFileUploaderDropzoneInstructions"]{color:var(--tx3)!important; font-size:0.82rem!important;}

[data-testid="stFormSubmitButton"] button{
  background:linear-gradient(135deg,#8A6A1F,#C9A84C)!important; color:#07080F!important;
  border:none!important; border-radius:6px!important; font-family:'Outfit',sans-serif!important;
  font-weight:600!important; font-size:0.82rem!important; letter-spacing:0.1em!important;
  text-transform:uppercase!important; padding:0.65rem 1.5rem!important;
  box-shadow:0 2px 12px rgba(201,168,76,0.2)!important; transition:all 0.2s!important;
}
[data-testid="stFormSubmitButton"] button:hover{
  background:linear-gradient(135deg,#C9A84C,#E4C06E)!important;
  box-shadow:0 4px 20px rgba(201,168,76,0.35)!important; transform:translateY(-1px)!important;
}
div.stButton>button{
  background:var(--p3)!important; color:var(--tx2)!important; border:1px solid var(--b2)!important;
  border-radius:6px!important; font-family:'Outfit',sans-serif!important;
  font-weight:500!important; font-size:0.8rem!important; padding:0.5rem 1rem!important;
  transition:all 0.2s!important;
}
div.stButton>button:hover{
  border-color:var(--gold-d)!important; color:var(--gold-b)!important;
  background:rgba(201,168,76,0.06)!important;
}
[data-testid="stDownloadButton"] button{
  background:transparent!important; color:var(--gold)!important;
  border:1px solid rgba(201,168,76,0.35)!important; border-radius:6px!important;
  font-family:'Outfit',sans-serif!important; font-weight:500!important; font-size:0.8rem!important;
  transition:all 0.2s!important;
}
[data-testid="stDownloadButton"] button:hover{
  background:rgba(201,168,76,0.08)!important; border-color:var(--gold)!important;
  box-shadow:0 2px 12px rgba(201,168,76,0.15)!important; transform:translateY(-1px)!important;
}
[data-testid="stDataFrame"]{border:1px solid var(--b2)!important; border-radius:8px!important; overflow:hidden!important;}
[data-testid="stDataFrame"] th{
  background:var(--p3)!important; color:var(--tx3)!important; font-size:0.7rem!important;
  text-transform:uppercase!important; letter-spacing:0.1em!important;
  border-bottom:1px solid var(--b2)!important;
}
[data-testid="stDataFrame"] td{color:var(--tx2)!important; font-size:0.82rem!important; border-bottom:1px solid var(--b1)!important;}
[data-testid="stExpander"]{background:var(--p2)!important; border:1px solid var(--b1)!important; border-radius:8px!important;}
.stSpinner>div{border-top-color:var(--gold)!important;}
::-webkit-scrollbar{width:5px;height:5px;}
::-webkit-scrollbar-track{background:var(--navy);}
::-webkit-scrollbar-thumb{background:var(--b2);border-radius:3px;}
::-webkit-scrollbar-thumb:hover{background:var(--gold-d);}
#MainMenu,footer,[data-testid="stToolbar"],[data-testid="stDecoration"]{display:none!important;}
.block-container{padding:2rem 2.5rem!important;}
</style>
"""

# ═══════════════════════════════════════════════
# COMPONENTES VISUAIS
# ═══════════════════════════════════════════════

@st.cache_data(show_spinner=False)
def _logo_b64(path: str) -> Optional[str]:
    if not os.path.exists(path): return None
    with open(path, "rb") as f: return base64.b64encode(f.read()).decode()

def secao(icone: str, titulo: str, sub: str = ""):
    linha_sub = (f'<div style="color:#555A70;font-size:0.72rem;letter-spacing:0.1em;'
                 f'text-transform:uppercase;margin-top:3px;">{sub}</div>') if sub else ""
    st.markdown(
        f'<div style="margin:1.5rem 0 1.8rem;">'
        f'<div style="display:flex;align-items:center;gap:14px;">'
        f'<div style="width:36px;height:36px;background:rgba(201,168,76,0.1);'
        f'border:1px solid rgba(201,168,76,0.25);border-radius:8px;'
        f'display:flex;align-items:center;justify-content:center;font-size:1rem;">{icone}</div>'
        f'<div><div style="font-family:\'Cormorant Garamond\',serif;font-size:1.5rem;'
        f'font-weight:700;color:#E8E4D9;line-height:1.1;">{titulo}</div>{linha_sub}</div>'
        f'</div>'
        f'<div style="height:1px;background:linear-gradient(90deg,rgba(201,168,76,0.4),'
        f'rgba(201,168,76,0.05),transparent);margin-top:16px;"></div></div>',
        unsafe_allow_html=True)

def metrica(label: str, valor: str, cor: str = "#C9A84C"):
    st.markdown(
        f'<div style="background:#0D0F1A;border:1px solid #1E2236;border-radius:10px;'
        f'padding:1.2rem 1.4rem;text-align:center;">'
        f'<div style="color:{cor};font-family:\'Cormorant Garamond\',serif;'
        f'font-size:2rem;font-weight:700;line-height:1;">{valor}</div>'
        f'<div style="color:#555A70;font-size:0.68rem;letter-spacing:0.12em;'
        f'text-transform:uppercase;margin-top:6px;">{label}</div></div>',
        unsafe_allow_html=True)

def card_processo(row: dict):
    try: dt = pd.to_datetime(row["data_envio"]).strftime("%d/%m/%Y %H:%M")
    except: dt = row.get("data_envio") or "—"
    st.markdown(
        f'<div style="background:#0D0F1A;border:1px solid #1E2236;border-left:3px solid #C9A84C;'
        f'border-radius:0 10px 10px 0;padding:1.1rem 1.4rem;margin-bottom:0.5rem;">'
        f'<div style="display:flex;justify-content:space-between;align-items:center;">'
        f'<div><div style="font-family:\'Cormorant Garamond\',serif;font-size:1.05rem;'
        f'font-weight:600;color:#E8E4D9;">{row["nome_cliente"]}</div>'
        f'<div style="font-family:\'JetBrains Mono\',monospace;color:#C9A84C;font-size:0.78rem;'
        f'margin-top:3px;">{row["numero_processo"]}</div></div>'
        f'<div style="text-align:right;">'
        f'<div style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);'
        f'color:#C9A84C;font-size:0.68rem;letter-spacing:0.08em;text-transform:uppercase;'
        f'padding:2px 10px;border-radius:20px;margin-bottom:6px;">{row["tipo"]}</div>'
        f'<div style="color:#555A70;font-size:0.7rem;">{dt}</div>'
        f'</div></div></div>', unsafe_allow_html=True)

def card_finalizado(row: dict):
    try: dt = pd.to_datetime(row["data_envio"]).strftime("%d/%m/%Y %H:%M")
    except: dt = row.get("data_envio") or "—"
    st.markdown(
        f'<div style="background:#0D0F1A;border:1px solid #1E2236;border-left:3px solid #2E7D5E;'
        f'border-radius:0 10px 10px 0;padding:1.1rem 1.4rem;">'
        f'<div style="display:flex;justify-content:space-between;align-items:center;">'
        f'<div><div style="font-family:\'Cormorant Garamond\',serif;font-size:1.05rem;'
        f'font-weight:600;color:#E8E4D9;">{row["nome_cliente"]}</div>'
        f'<div style="font-family:\'JetBrains Mono\',monospace;color:#3DAA7B;font-size:0.78rem;'
        f'margin-top:3px;">{row["numero_processo"]}</div></div>'
        f'<div style="color:#555A70;font-size:0.7rem;">{dt}</div>'
        f'</div></div>', unsafe_allow_html=True)

def vazio(icone: str, msg: str):
    st.markdown(
        f'<div style="text-align:center;padding:4rem 2rem;">'
        f'<div style="font-size:2.5rem;margin-bottom:1rem;opacity:0.3;">{icone}</div>'
        f'<div style="font-family:\'Cormorant Garamond\',serif;font-size:1.1rem;'
        f'color:#555A70;">{msg}</div></div>', unsafe_allow_html=True)

def sidebar_logo():
    logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
    enc = _logo_b64(logo_path)
    img = (f'<img src="data:image/png;base64,{enc}" style="width:70px;margin-bottom:12px;"/>'
           if enc else '<div style="font-size:2.5rem;margin-bottom:12px;">⚖</div>')
    st.sidebar.markdown(
        f'<div style="padding:2rem 0 1.5rem;text-align:center;border-bottom:1px solid #1E2236;">'
        f'{img}'
        f'<div style="font-family:\'Cormorant Garamond\',serif;font-size:1.4rem;font-weight:700;'
        f'color:#C9A84C;letter-spacing:0.18em;">JUSREPORT</div>'
        f'<div style="font-size:0.6rem;color:#555A70;letter-spacing:0.25em;text-transform:uppercase;'
        f'margin-top:4px;">Legal Due Diligence</div></div>', unsafe_allow_html=True)

def sidebar_status(gem_ok: bool):
    st.sidebar.markdown(
        f'<div style="padding:1.2rem 1.4rem;border-top:1px solid #1E2236;margin-top:auto;">'
        f'<div style="font-size:0.62rem;color:#555A70;letter-spacing:0.15em;'
        f'text-transform:uppercase;margin-bottom:10px;">Sistema</div>'
        f'<div style="display:flex;justify-content:space-between;">'
        f'<span style="font-size:0.75rem;color:#9B97A6;">Gemini</span>'
        f'<span style="font-size:0.75rem;color:{"#3DAA7B" if gem_ok else "#E57373"};font-weight:600;">'
        f'{"● Ativo" if gem_ok else "● Inativo"}</span></div></div>',
        unsafe_allow_html=True)

# ═══════════════════════════════════════════════
# APP
# ═══════════════════════════════════════════════

st.set_page_config(page_title="JusReport", page_icon="⚖️",
                   layout="wide", initial_sidebar_state="expanded")
st.markdown(CSS, unsafe_allow_html=True)

gem_ok = bool(GEMINI_API_KEY)

sidebar_logo()
st.sidebar.markdown(
    '<div style="padding:1rem 1.2rem 0.5rem;">'
    '<div style="font-size:0.62rem;color:#555A70;letter-spacing:0.15em;'
    'text-transform:uppercase;margin-bottom:8px;">Navegação</div></div>',
    unsafe_allow_html=True)

pagina = st.sidebar.radio("", [
    "📂  Novo Processo",
    "⏳  Pendentes",
    "✅  Finalizados",
    "📅  Relatório Mensal",
], label_visibility="collapsed")

sidebar_status(gem_ok)

if not gem_ok:
    st.warning("⚠ GEMINI_API_KEY não configurada — processamento automático desativado.")

# ══════════════════════════════════════════
# NOVO PROCESSO
# ══════════════════════════════════════════
if "Novo" in pagina:
    secao("📂", "Novo Processo", "Cadastre um processo para análise automática")

    df_p = pendentes_df(); df_f = finalizados_df()
    c1, c2, c3 = st.columns(3)
    with c1: metrica("Pendentes",   str(len(df_p)), "#C9A84C")
    with c2: metrica("Finalizados", str(len(df_f)), "#3DAA7B")
    with c3: metrica("Total",       str(len(df_p) + len(df_f)), "#9B97A6")

    st.markdown('<div style="margin-top:1.5rem;"></div>', unsafe_allow_html=True)

    with st.form("form_novo"):
        col_a, col_b = st.columns(2)
        with col_a:
            colaborador = st.text_input("Colaborador")
            numero      = st.text_input("Número do processo")
        with col_b:
            tipo    = st.selectbox("Tipo de sumarização", SUMARIZACOES)
            arquivo = st.file_uploader("Arquivo do processo", type=["pdf", "docx"])

        st.markdown('<div style="margin-top:1rem;"></div>', unsafe_allow_html=True)
        enviado = st.form_submit_button("Cadastrar Processo", use_container_width=True)

        if enviado:
            if not (colaborador and numero and arquivo):
                st.warning("Preencha todos os campos antes de enviar.")
            else:
                try:
                    salvar_processo(nome_cliente=colaborador, email="", numero=numero,
                                    tipo=tipo, arquivo_bytes=arquivo.getvalue(),
                                    nome_arquivo=arquivo.name, conferencia="interno")
                    st.success("✅ Processo cadastrado com sucesso!")
                    invalidar()
                except Exception as e:
                    st.error(f"Erro ao cadastrar: {e}")
                    with st.expander("Detalhes"):
                        st.code("".join(traceback.format_exception(type(e), e, e.__traceback__)))

# ══════════════════════════════════════════
# PENDENTES
# ══════════════════════════════════════════
elif "Pendentes" in pagina:
    df = pendentes_df()
    secao("⏳", "Processos Pendentes",
          f"{len(df)} processo{'s' if len(df) != 1 else ''} aguardando")

    if df.empty:
        vazio("⚖", "Nenhum processo pendente no momento")
    else:
        for _, row in df.iterrows():
            card_processo(dict(row))
            col1, col2, col3 = st.columns([3, 1, 1])

            with col1:
                caminho = row.get("caminho_arquivo")
                if caminho and os.path.exists(caminho):
                    with open(caminho, "rb") as f:
                        st.download_button("↓  Baixar original", data=f,
                            file_name=os.path.basename(caminho),
                            mime=_guess_mime(os.path.basename(caminho)),
                            key=f"dl_orig_{row['id']}", use_container_width=True)

            with col2:
                if st.button("Excluir", key=f"del_{row['id']}", use_container_width=True):
                    try:
                        excluir_com_arquivo(row["id"], row.get("caminho_arquivo"))
                        st.success("Excluído."); invalidar(); st.rerun()
                    except Exception as e:
                        st.error(f"Erro: {e}")

            with col3:
                if not gem_ok:
                    st.button("Processar", key=f"proc_{row['id']}",
                              disabled=True, use_container_width=True)
                elif st.button("Processar", key=f"proc_{row['id']}",
                               use_container_width=True):
                    caminho = row.get("caminho_arquivo")
                    if not caminho or not os.path.exists(caminho):
                        st.error("Arquivo não encontrado.")
                    else:
                        try:
                            pbar = st.progress(0, text="Extraindo texto do PDF...")
                            base_text = extrair_texto_pdf(caminho)
                            if not base_text:
                                st.error("Não foi possível extrair texto do PDF.")
                                st.stop()

                            pbar.progress(10, text="Iniciando análise com IA...")
                            summary_md = gerar_relatorio_md(
                                base_text=base_text,
                                case_number=str(row["numero_processo"]),
                                action_type=str(row["tipo"]),
                                progress_bar=pbar,
                            )
                            pbar.progress(95, text="Gerando documento DOCX...")
                            docx_bytes = gerar_docx(summary_md)
                            pbar.progress(100, text="Concluído!")
                            time.sleep(0.3)
                            pbar.empty()

                            nome_saida    = f"JusReport_{row['numero_processo']}.docx"
                            caminho_rel   = os.path.join(RELATORIOS_DIR, nome_saida)
                            with open(caminho_rel, "wb") as out:
                                out.write(docx_bytes)

                            registrar_relatorio(row["id"], caminho_docx=caminho_rel)
                            invalidar()

                            st.markdown(
                                '<div style="background:rgba(46,125,82,0.1);'
                                'border:1px solid rgba(61,170,123,0.3);border-radius:8px;'
                                'padding:1rem 1.4rem;margin:0.8rem 0;">'
                                '<span style="color:#3DAA7B;font-weight:600;">✓</span>'
                                '<span style="color:#E8E4D9;margin-left:8px;">'
                                'Relatório gerado com sucesso</span></div>',
                                unsafe_allow_html=True)

                            st.download_button(
                                "↓  Baixar Relatório", data=docx_bytes,
                                file_name=nome_saida,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_novo_{row['id']}", use_container_width=True)

                        except Exception as e:
                            st.error(f"Erro no processamento: {e}")
                            with st.expander("Detalhes técnicos"):
                                st.code("".join(traceback.format_exception(type(e), e, e.__traceback__)))

            st.markdown('<div style="margin-bottom:1rem;"></div>', unsafe_allow_html=True)

# ══════════════════════════════════════════
# FINALIZADOS
# ══════════════════════════════════════════
elif "Finalizados" in pagina:
    df = finalizados_df()
    secao("✅", "Relatórios Finalizados",
          f"{len(df)} documento{'s' if len(df) != 1 else ''} disponíve{'is' if len(df) != 1 else 'l'}")

    if df.empty:
        vazio("📄", "Nenhum relatório finalizado ainda")
    else:
        try:
            df["data_envio"] = pd.to_datetime(df["data_envio"]).dt.strftime("%d/%m/%Y %H:%M")
        except: pass

        for _, row in df.iterrows():
            col_info, col_dl = st.columns([5, 1])
            with col_info: card_finalizado(dict(row))
            with col_dl:
                caminho_rel = row.get("caminho_relatorio")
                if caminho_rel and os.path.exists(caminho_rel):
                    with open(caminho_rel, "rb") as f:
                        st.download_button("↓  Baixar", data=f,
                            file_name=os.path.basename(caminho_rel),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_fin_{row['id']}", use_container_width=True)
                else:
                    st.caption("Não encontrado")
            st.markdown('<div style="margin-bottom:0.3rem;"></div>', unsafe_allow_html=True)

        st.markdown('<div style="margin-top:1.5rem;"></div>', unsafe_allow_html=True)
        out = BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as w:
            df[["nome_cliente","numero_processo","data_envio"]].to_excel(
                w, index=False, sheet_name="Finalizados")
        st.download_button("↓  Exportar lista em Excel", data=out.getvalue(),
            file_name="relatorios_finalizados.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ══════════════════════════════════════════
# RELATÓRIO MENSAL
# ══════════════════════════════════════════
elif "Mensal" in pagina:
    secao("📅", "Relatório Mensal", "Volume de processos por colaborador")

    df = mensal_df()
    if df.empty:
        vazio("📊", "Nenhum dado disponível ainda")
    else:
        st.dataframe(df, use_container_width=True, hide_index=True)
        out = BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as w:
            df.to_excel(w, index=False, sheet_name="Mensal")
        st.markdown('<div style="margin-top:1rem;"></div>', unsafe_allow_html=True)
        st.download_button("↓  Exportar em Excel", data=out.getvalue(),
            file_name="relatorio_mensal.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

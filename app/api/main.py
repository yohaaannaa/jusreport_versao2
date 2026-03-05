import os
import uuid
import io
import traceback
from typing import Dict, Any

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
import pdfplumber
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from pydantic import BaseModel

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("[AVISO] pdf2image não está instalado. Prints de planilhas não serão gerados.")


# ============================================================
# CONFIGURAÇÃO (PATHS, .ENV, GEMINI)
# ============================================================

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
load_dotenv(os.path.join(BASE_DIR, ".env"))

DATA_DIR = os.path.join(BASE_DIR, "data")
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
REL_DIR = os.path.join(DATA_DIR, "relatorios")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REL_DIR, exist_ok=True)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL_TEXT = os.getenv("GEMINI_MODEL_TEXT", "gemini-2.5-pro")

# Origens permitidas (sobrescreva via .env em produção)
ALLOWED_ORIGINS = os.getenv("CORS_ALLOWED_ORIGINS", "*").split(",")

if GEMINI_API_KEY:
    print(f"[INFO] GEMINI_API_KEY detectada (prefixo={GEMINI_API_KEY[:6]}...)")
    genai.configure(api_key=GEMINI_API_KEY)
    try:
        text_model = genai.GenerativeModel(GEMINI_MODEL_TEXT)
        print(f"[INFO] Carregado modelo Gemini: {GEMINI_MODEL_TEXT}")
    except Exception as e:
        print(f"[AVISO] Falha ao carregar {GEMINI_MODEL_TEXT}: {e}. Tentando fallback...")
        GEMINI_MODEL_TEXT = "gemini-1.5-pro"
        text_model = genai.GenerativeModel(GEMINI_MODEL_TEXT)
        print(f"[INFO] Usando fallback: {GEMINI_MODEL_TEXT}")
else:
    print("[AVISO] GEMINI_API_KEY não configurada. IA desativada.")
    text_model = None


# ============================================================
# PALAVRAS-CHAVE (única fonte da verdade — sem duplicação)
# ============================================================

HOTSPOT_KEYWORDS = [
    "planilha",
    "demonstrativo",
    "cálculo",
    "calculo",
    "sisbajud",
    "bacenjud",
    "bloqueio",
    "penhora online",
    "penhora on-line",
]


# ============================================================
# FASTAPI + CORS
# ============================================================

app = FastAPI(title="API Jurídica - JusReport")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Jobs em memória (futuro: substituir por Redis ou banco)
JOBS: Dict[str, Dict[str, Any]] = {}

# Limite de jobs em memória para evitar vazamento
MAX_JOBS = 500


# ============================================================
# MODELO P/ CORPO DO /summarize (JSON)
# ============================================================

class SummarizeRequest(BaseModel):
    question: str
    case_number: str
    action_type: str
    k: int = 50
    return_json: bool = True


# ============================================================
# ENDPOINTS BÁSICOS
# ============================================================

@app.get("/health")
def health():
    env_val = os.getenv("GEMINI_API_KEY")
    return {
        "service": "api-juridica",
        "gemini_env_present": env_val is not None,
        "gemini_env_prefix": (env_val[:6] + "...") if env_val else None,
        "gemini_configured": bool(env_val),
        "gemini_model": GEMINI_MODEL_TEXT if env_val else None,
    }


@app.post("/ingest")
async def ingest(
    files: list[UploadFile] = File(...),
    case_number: str = Form(...),
    client_id: str | None = Form(None),
):
    if not files:
        raise HTTPException(status_code=400, detail="Nenhum arquivo enviado")

    # Limita tamanho do dicionário para evitar vazamento de memória
    if len(JOBS) >= MAX_JOBS:
        oldest_key = next(iter(JOBS))
        del JOBS[oldest_key]

    f = files[0]
    job_id = str(uuid.uuid4())
    filename = f"{job_id}__{f.filename}"
    save_path = os.path.join(UPLOAD_DIR, filename)

    content = await f.read()
    with open(save_path, "wb") as out:
        out.write(content)

    JOBS[job_id] = {
        "status": "done",
        "progress": 100,
        "detail": "Ingestão concluída",
        "file_path": save_path,
        "case_number": case_number,
        "client_id": client_id,
        "meta": {},
    }

    return {"job_id": job_id}


@app.get("/status/{job_id}")
def status(job_id: str):
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job não encontrado")
    return {
        "status": job["status"],
        "progress": job["progress"],
        "detail": job.get("detail", ""),
        "result": None,
    }


# ============================================================
# EXTRAÇÃO DE TEXTO DO PDF
# ============================================================

def _extract_text_from_pdf(path: str) -> tuple[str, Dict[str, Any]]:
    text_by_page: list[str] = []
    pages_obj = []

    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages_obj.append(page)
                text_by_page.append(page.extract_text() or "")
    except Exception as e:
        print(f"[ERRO] Falha ao ler PDF {path}: {e}")

    full_text = "\n\n".join(text_by_page)
    total_len = len(full_text)

    if total_len == 0:
        return "", {"planilha_pages": []}

    env_max = int(os.getenv("MAX_PDF_CHARS", "30000"))
    HARD_CAP_CHARS = 80000
    max_chars = min(env_max, HARD_CAP_CHARS)

    print(
        f"[INFO] MAX_PDF_CHARS={env_max} | HARD_CAP={HARD_CAP_CHARS} "
        f"| efetivo={max_chars} | total={total_len}"
    )

    if total_len <= max_chars:
        planilha_pages = _detect_planilha_pages(text_by_page)
        return full_text, {"planilha_pages": planilha_pages}

    # Hotspots
    hotspot_parts: list[str] = []
    planilha_pages: list[int] = []

    for idx, page_text in enumerate(text_by_page):
        raw = page_text or ""
        tl = raw.lower()
        if any(k in tl for k in HOTSPOT_KEYWORDS):
            page_num = idx + 1
            planilha_pages.append(page_num)
            bloco = [f"\n\n=== PÁGINA RELEVANTE {page_num} ===\n\n", raw]

            try:
                tables = pages_obj[idx].extract_tables() or []
            except Exception as te:
                print(f"[AVISO] Falha ao extrair tabelas da pág. {page_num}: {te}")
                tables = []

            if tables:
                bloco.append(f"\n\n=== TABELA(S) NA PÁGINA {page_num} ===\n\n")
                for t_idx, table in enumerate(tables, 1):
                    bloco.append(f"--- Tabela {t_idx} ---\n")
                    for row in table:
                        linha = " | ".join(cell or "" for cell in row)
                        bloco.append(linha + "\n")
                    bloco.append("\n")

            hotspot_parts.append("".join(bloco))

    hotspot_text = "".join(hotspot_parts).strip()

    if not hotspot_text:
        return _build_global_sample(full_text, max_chars), {"planilha_pages": []}

    max_hotspot = int(max_chars * 0.6)
    if len(hotspot_text) > max_hotspot:
        hotspot_text = hotspot_text[:max_hotspot]

    remaining_chars = max_chars - len(hotspot_text)
    if remaining_chars <= 0:
        return hotspot_text, {"planilha_pages": planilha_pages}

    global_sample = _build_global_sample(full_text, remaining_chars)
    final_text = hotspot_text + "\n\n=== AMOSTRAGEM GLOBAL ===\n\n" + global_sample

    print(f"[INFO] Texto final: {len(final_text)} chars (hotspots={len(hotspot_text)}, global={len(global_sample)})")
    return final_text, {"planilha_pages": planilha_pages}


def _detect_planilha_pages(text_by_page: list[str]) -> list[int]:
    return [
        idx + 1
        for idx, text in enumerate(text_by_page)
        if any(k in (text or "").lower() for k in HOTSPOT_KEYWORDS)
    ]


def _build_global_sample(full_text: str, max_chars: int) -> str:
    total_len = len(full_text)
    if total_len <= max_chars:
        return full_text

    part = max(max_chars // 4, 1)
    inicio = full_text[:part]

    mid_center = total_len // 2
    mid_start = max(0, mid_center - part // 2)
    meio = full_text[mid_start: mid_start + part]

    pre_final_start = max(0, total_len - part * 2)
    pre_final = full_text[pre_final_start: pre_final_start + part]

    fim = full_text[-part:]

    return (
        inicio
        + "\n\n=== TRECHO CENTRAL ===\n\n" + meio
        + "\n\n=== TRECHO PRÉ-FINAL ===\n\n" + pre_final
        + "\n\n=== TRECHO FINAL ===\n\n" + fim
    )


# ============================================================
# AGENTES DE SUMARIZAÇÃO (chamadas paralelas ao Gemini)
# ============================================================

def _build_tasks(action_type: str) -> list[dict]:
    """
    9 seções do Relatório de Pré-Auditoria / Legal Due Diligence NPL.
    Cada agente é responsável por uma seção do relatório final.
    """
    return [
        # ------------------------------------------------------------------
        # SEÇÃO 1 — SUMÁRIO EXECUTIVO
        # ------------------------------------------------------------------
        {
            "key": "sumario_executivo",
            "title": "1. Sumário Executivo",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL (Non-Performing Loans).

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 1. SUMÁRIO EXECUTIVO" em Markdown.

Esta seção deve conter UMA TABELA com três colunas: **INDICADOR** | **DADOS EXTRAÍDOS** | **FONTE**

Preencha obrigatoriamente as seguintes linhas (se o dado não constar, escreva "Não localizado"):

| INDICADOR | DADOS EXTRAÍDOS | FONTE |
| :--- | :--- | :--- |
| Processo nº | | |
| Vara / Comarca | | |
| Data de Distribuição | | |
| Classe | | |
| Exequente (Credor) | nome completo + CNPJ/CPF | |
| Executado Principal | nome completo + CPF/CNPJ se constar | |
| Co-executados | nome(s) + qualificação | |
| Advogado Credor (atual) | nome + OAB | |
| Advogado(s) dos Executados | nome + OAB | |
| Valor da Causa (Inicial) | R$ + data | |
| Valor Atualizado (Exigível) | R$ + data base | |
| Título Executivo | tipo + número | |
| Bem Penhorado / Garantia Principal | descrição resumida + matrícula/nº | |
| Avaliação do Bem | R$ + data | |
| Risco de Prescrição | BAIXO / MÉDIO / ALTO + justificativa curta | |

Na coluna FONTE, indique sempre número do documento, página ou referência encontrada no texto.
Não invente dados. Se um campo não constar nos trechos analisados, escreva "Não localizado".
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 2 — MAPA DE CITAÇÃO E PARTES
        # ------------------------------------------------------------------
        {
            "key": "mapa_citacao",
            "title": "2. Mapa de Citação e Partes — Auditoria de Nulidades",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 2. MAPA DE CITAÇÃO E PARTES — AUDITORIA DE NULIDADES" em Markdown.

Para CADA parte do processo (executado principal, co-executados, fiadores, avalistas,
terceiros interessados que tenham ingressado nos autos), crie uma subseção "## 2.X Nome — Qualificação"
contendo os seguintes campos:

**Qualificação:** endereço, CPF/CNPJ se constar.
**Status da Citação:** Citado / Não citado / Citação não localizada nos autos digitais.
**Forma:** como foi citada (AR, edital, oficial de justiça, processo físico pré-digitalização etc.).
**Representação:** advogado constituído, se houver — nome e OAB.
**Evento Morte/Sucessão:** se houver notícia de falecimento ou habilitação de herdeiros; se não houver, escrever "Não localizado".
**Risco de Nulidade:** BAIXO / MÉDIO / ALTO — com justificativa de 1–3 linhas.

Se terceiros interessados ingressaram nos autos (embargos de terceiro, pedido de desconstituição etc.),
crie uma subseção específica para cada um, detalhando: qualificação, data de ingresso,
tese apresentada, resposta do credor e status atual.

Regras:
- Não invente dados; se não constar, escreva "Não localizado nos autos analisados".
- Mencione número do documento e página sempre que possível.
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 3 — TÍTULO E ORIGEM DA DÍVIDA
        # ------------------------------------------------------------------
        {
            "key": "titulo_origem",
            "title": "3. Título e Origem da Dívida",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 3. TÍTULO E ORIGEM DA DÍVIDA" em Markdown.

Esta seção deve conter UMA TABELA com três colunas: **ASPECTO** | **DESCRIÇÃO** | **FONTE**

Preencha obrigatoriamente as seguintes linhas:

| ASPECTO | DESCRIÇÃO | FONTE |
| :--- | :--- | :--- |
| Título Executivo | tipo (CCB, Cédula de Crédito Rural, Cheque, etc.) | |
| Identificação / Nº do título | número interno ou do documento | |
| Classe Processual | código e descrição | |
| Credor Original | nome + CNPJ/CPF | |
| Devedor Principal | nome + CPF/CNPJ | |
| Co-devedores / Avalistas | nomes | |
| Data de Distribuição | dd/mm/aaaa | |
| Garantia Original | tipo (hipoteca, penhor, aval, fiança) + descrição do bem/obrigação | |
| Valor Original da Causa | R$ | |
| Valor Atualizado | R$ + data base | |
| Registro de Imóveis | matrícula(s), CRI, cadeia dominial se constar | |
| Observações Relevantes | migração de físico para digital, renegociações, etc. | |

Após a tabela, se houver alguma NOTE CRÍTICA sobre a garantia (ex.: hipoteca cedular, fraude à execução
arguida, questão dominial), inclua em parágrafo destacado com o prefixo "*NOTA CRÍTICA — [TEMA]:*".

Não invente dados. Use "Não localizado" quando o dado não constar.
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 4 — RASTREAMENTO DE ATIVOS (ASSET TRACING)
        # ------------------------------------------------------------------
        {
            "key": "asset_tracing",
            "title": "4. Rastreamento de Ativos — Asset Tracing",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 4. RASTREAMENTO DE ATIVOS — ASSET TRACING" em Markdown, com as subseções abaixo:

## 4.1 Penhoras Efetivadas
Crie uma tabela com colunas **BEM** | **DESCRIÇÃO** | **FONTE** listando cada bem penhorado.
Para cada bem inclua: tipo (imóvel, veículo, dinheiro, etc.), descrição completa
(área, matrícula, localização, valor), data do termo de penhora, depositário fiel nomeado,
status do registro da penhora (averbado / pendente / não localizado) e referência documental.
Se não houver penhora efetivada, escreva "Não localizado".

## 4.2 Avaliação do Bem Penhorado
Crie uma tabela com colunas **ITEM** | **DADOS** | **FONTE** contendo:
avaliador, data da avaliação (vistoria), área/descrição avaliada, metodologia,
VALOR DA AVALIAÇÃO, status da homologação, impugnações apresentadas e resultado,
relação dívida/garantia (cobertura em "x"), uso atual do bem.
Se não houver avaliação, escreva "Não localizado".

## 4.3 Diligências de Busca de Bens (Sistemas)
Para cada sistema abaixo, informe se há pedido, decisão, resultado e referência:
- SISBAJUD / BACENJUD (bloqueio de ativos financeiros)
- RENAJUD (veículos)
- INFOJUD (declarações fiscais)
- SERASAJUD
- Outros (cartas precatórias, pesquisas em cartórios, CCS, INTERMAT, etc.)

Se não encontrar menção a um sistema, escreva:
"Não localizado nos trechos analisados — ausência de registros de [SISTEMA]."

Regras:
- Mencione Num. do documento e página sempre que possível.
- Não invente dados.
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 5 — STATUS EXPROPRIATÓRIO (LEILÃO / ADJUDICAÇÃO)
        # ------------------------------------------------------------------
        {
            "key": "status_expropriatorio",
            "title": "5. Status Expropriatório — Leilão / Adjudicação",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 5. STATUS EXPROPRIATÓRIO — LEILÃO / ADJUDICAÇÃO" em Markdown.

Crie uma tabela com colunas **DATA / ATO** | **DESCRIÇÃO** | **STATUS** | **FONTE**
listando em ordem cronológica TODOS os atos relacionados a:
- Pedidos de designação de leilão ou hasta pública
- Decisões sobre leilão (deferimento, indeferimento, suspensão, condicionamentos)
- Embargos de declaração sobre essas decisões
- Incidentes que bloqueiem ou retardem o leilão
- Designação efetiva de leilão (1ª e 2ª praças), com datas e valores mínimos
- Arrematação ou adjudicação, se houver
- Status atual (suspenso, pendente, aguardando decisão, designado para [data], etc.)

Após a tabela, escreva um parágrafo de **CONCLUSÃO** em negrito resumindo o estágio atual
do processo expropriatório e o principal obstáculo, se houver.

Não invente dados. Use "Não designado" ou "Não realizado" conforme o caso.
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 6 — DEFESAS E INCIDENTES
        # ------------------------------------------------------------------
        {
            "key": "defesas_incidentes",
            "title": "6. Defesas e Incidentes — O Passivo",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 6. DEFESAS E INCIDENTES — O PASSIVO" em Markdown.

Para CADA defesa ou incidente processual encontrado, crie uma subseção numerada
"## 6.X [Nome do Incidente / Defesa]" contendo:

**Data:** data da apresentação ou decisão.
**Documento:** número do documento (Num.) ou folha (fls.).
**Requerente / Autor:** quem apresentou.
**Tese Principal:** resumo em 2–5 linhas da tese ou pedido.
**Decisão:** resultado e fundamento legal, com data.
**Recurso / Desdobramento:** se houve recurso, embargos de declaração, agravo etc., descreva e informe o resultado.
**Status:** REJEITADO DEFINITIVAMENTE / PENDENTE DE DECISÃO / ACOLHIDO / EM RECURSO.

Inclua TODOS os incidentes que encontrar, especialmente:
- Embargos à execução (Lei 6.830/80 ou art. 914 CPC)
- Exceção de pré-executividade
- Impugnação à avaliação / substituição de penhora
- Embargos de terceiro / pedido de desconstituição de penhora
- Incidente de fraude à execução
- Incidente de desconsideração da personalidade jurídica
- Embargos de declaração relevantes
- Prescrição intercorrente (se arguida)

Se não localizar determinada espécie de defesa, crie subseção com o título e escreva:
"Não localizado nos autos digitalizados analisados."
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 7 — RED FLAGS E RISCOS CRÍTICOS
        # ------------------------------------------------------------------
        {
            "key": "red_flags",
            "title": "7. Red Flags e Riscos Críticos",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 7. RED FLAGS E RISCOS CRÍTICOS" em Markdown.

Esta seção deve conter dois grupos:

**GRUPO 1 — RED FLAGS (riscos que podem comprometer a recuperação do crédito)**
Para cada red flag identificado, use o formato:
**⚠ RED FLAG #N — [TÍTULO DO RISCO] [RISCO ALTO / MÉDIO]**
Parágrafo descrevendo: o que foi encontrado, por que representa risco, qual o impacto potencial
na execução ou recuperação do crédito, e número(s) de documento de referência.

Exemplos de red flags a pesquisar:
- Disputa de propriedade do bem penhorado (terceiro interessado)
- Penhora não averbada no registro de imóveis
- Citação de executado não comprovada nos autos digitais
- Alteração de matrícula do imóvel penhorado
- Risco de prescrição intercorrente
- Processo físico com possíveis peças não digitalizadas
- Outros riscos que você identificar

**GRUPO 2 — PONTOS POSITIVOS (fatores favoráveis à recuperação)**
Para cada ponto positivo, use o formato:
**✔ PONTO POSITIVO #N — [TÍTULO]**
Parágrafo descrevendo o fator e seu impacto positivo na análise.

Exemplos: cobertura patrimonial excepcional, prescrição descartada, decisões favoráveis transitadas,
tese jurídica sólida do credor etc.

Regras:
- Seja objetivo e factual. Não exagere nem minimize os riscos.
- Fundamente cada item com referência documental (Num., fls., data).
- Se não encontrar red flags, escreva "Nenhum red flag crítico identificado nos trechos analisados."
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 8 — ÍNDICE REMISSIVO / LINHA DO TEMPO
        # ------------------------------------------------------------------
        {
            "key": "linha_do_tempo",
            "title": "8. Índice Remissivo — Linha do Tempo",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 8. ÍNDICE REMISSIVO — LINHA DO TEMPO" em Markdown.

Crie uma tabela com colunas **DATA** | **EVENTO** | **DOCUMENTO / FONTE**
listando EM ORDEM CRONOLÓGICA todos os atos processuais relevantes que encontrar, incluindo:

- Distribuição da ação
- Citações (data, forma, resultado)
- Migrações (processo físico → digital)
- Juntada de documentos relevantes (contratos, planilhas, laudos)
- Pedidos do credor e das partes
- Despachos e decisões interlocutórias
- Penhoras, avaliações, hasta pública
- Substabelecimentos de procuração
- Embargos, incidentes e seus desdobramentos
- Último ato registrado

Para datas inexatas, use "Mês/Ano" ou "Ano" com nota "(data aproximada)".
Na coluna DOCUMENTO / FONTE, inclua Num. do documento PJe ou referência de página.

Prefira listar MAIS atos do que menos. O objetivo é dar ao leitor uma visão completa
do histórico processual.
""",
        },

        # ------------------------------------------------------------------
        # SEÇÃO 9 — PARECER FINAL E RECOMENDAÇÃO
        # ------------------------------------------------------------------
        {
            "key": "parecer_final",
            "title": "9. Parecer Final e Recomendação de Aquisição",
            "instruction": """\
Você é um especialista em Legal Due Diligence de carteiras NPL.

Com base EXCLUSIVAMENTE no texto do processo abaixo, produza a seção
"# 9. PARECER FINAL E RECOMENDAÇÃO DE AQUISIÇÃO" em Markdown.

Esta seção deve ter a seguinte estrutura:

**SÍNTESE DA ANÁLISE:**
Parágrafo introdutório de 3–5 linhas resumindo a análise geral do crédito.

Em seguida, uma tabela de avaliação com colunas **CRITÉRIO** | **AVALIAÇÃO** | **COMENTÁRIO**:

| CRITÉRIO | AVALIAÇÃO | COMENTÁRIO |
| :--- | :--- | :--- |
| LIQUIDEZ | ALTA / MÉDIA-ALTA / MÉDIA / BAIXA | Descreva o ativo de garantia e estágio do processo |
| CERTEZA | ALTA / MÉDIA / BAIXA | Solidez do título executivo e das decisões favoráveis |
| EXIGIBILIDADE | ALTA / MÉDIA / BAIXA | Status do valor atualizado e eventuais questionamentos |
| SOLVÊNCIA DO DEVEDOR | AVALIAÇÃO PARCIAL / POSITIVA / NEGATIVA | Bens identificados, pesquisas realizadas |
| GARANTIA | EXISTENTE / EM RISCO / INEXISTENTE | Bem penhorado, valor, riscos sobre a garantia |

Após a tabela, escreva:

**RECOMENDAÇÃO: [AQUISIÇÃO RECOMENDADA / AQUISIÇÃO CONDICIONADA / AQUISIÇÃO NÃO RECOMENDADA]**

Se "CONDICIONADA", liste as condicionantes no formato:
- CONDIÇÃO 1 — [TÍTULO]: descrição da diligência ou requisito.
- CONDIÇÃO 2 — [TÍTULO]: ...

Finalize com um parágrafo de fechamento em itálico destacando o potencial de recuperação
e os principais riscos residuais.

Regras:
- Baseie-se APENAS no que consta nos autos. Não invente dados ou premissas.
- Seja objetivo e profissional. Este parecer é para tomada de decisão de investimento.
- Não omita riscos relevantes identificados na análise.
""",
        },
    ]


def _call_gemini(task: dict, base_text: str) -> tuple[str, str]:
    """Chama o Gemini para uma sub-tarefa. Retorna (key, texto)."""
    prompt = f"""
{task["instruction"]}

=== TEXTO DO PROCESSO (EXTRAÍDO DO PDF) ===

\"\"\"{base_text}\"\"\"
"""
    print(f"[AGENTE] Iniciando: {task['key']} ({task['title']})")
    resp = text_model.generate_content(prompt)
    text = (resp.text or "").strip()
    print(f"[AGENTE] Concluído: {task['key']} ({len(text)} chars)")
    return task["key"], text


def _run_execucao_agents(
    base_text: str, case_number: str, action_type: str
) -> tuple[str, dict]:
    """
    Envia sub-tarefas ao Gemini SEQUENCIALMENTE para economizar memória
    no plano gratuito do Render (512MB RAM).
    """
    if not text_model:
        raise RuntimeError("Modelo Gemini não configurado (text_model=None).")

    tasks = _build_tasks(action_type)
    sections: dict[str, str] = {}

    for task in tasks:
        try:
            key, text = _call_gemini(task, base_text)
            sections[key] = text
        except Exception as e:
            print(f"[ERRO] Sub-tarefa '{task['key']}' falhou: {e}")
            sections[task["key"]] = f"Erro ao processar esta seção: {e}"

    from datetime import date
    data_geracao = date.today().strftime("%d de %B de %Y").replace(
        "January","janeiro").replace("February","fevereiro").replace(
        "March","março").replace("April","abril").replace("May","maio").replace(
        "June","junho").replace("July","julho").replace("August","agosto").replace(
        "September","setembro").replace("October","outubro").replace(
        "November","novembro").replace("December","dezembro")

    header = (
        f"**RELATÓRIO DE PRÉ-AUDITORIA**\n\n"
        f"**LEGAL DUE DILIGENCE — NPL**\n\n"
        f"{action_type} — Processo nº {case_number}\n\n"
        f"*Data de Geração: {data_geracao} | CONFIDENCIAL*"
    )

    order = [t["key"] for t in tasks]
    md_parts = [header]
    for key in order:
        section_text = sections.get(key, "").strip()
        if section_text:
            md_parts.append(section_text)
        else:
            title = next(t["title"] for t in tasks if t["key"] == key)
            md_parts.append(f"# {title}\n\nNão localizado nos trechos analisados.")

    return "\n\n---\n\n".join(md_parts), sections


# ============================================================
# ENDPOINT /summarize
# ============================================================

@app.post("/summarize")
async def summarize(req: SummarizeRequest):
    try:
        job = next(
            (j for j in JOBS.values() if j.get("case_number") == req.case_number),
            None,
        )
        if not job:
            raise HTTPException(
                status_code=404,
                detail="Nenhum job encontrado para esse número de processo",
            )

        if not GEMINI_API_KEY or not text_model:
            raise HTTPException(
                status_code=500,
                detail="Gemini não configurado na API (.env)",
            )

        base_text, meta = _extract_text_from_pdf(job["file_path"])
        if not base_text:
            base_text = "Não foi possível extrair texto do PDF. Verifique o arquivo."

        job.setdefault("meta", {}).update(meta)

        final_md, sections = _run_execucao_agents(
            base_text=base_text,
            case_number=req.case_number,
            action_type=req.action_type,
        )

        return {
            "summary_markdown": final_md,
            "sections": sections,
            "used_chunks": [],
            "result": {"meta": meta},
        }

    except HTTPException:
        raise
    except Exception as e:
        print("ERRO EM /summarize:\n", traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"{e.__class__.__name__}: {e}")


# ============================================================
# ENDPOINT /export/docx
# ============================================================

@app.post("/export/docx")
async def export_docx(
    content: str = Form(...),
    filename: str = Form("relatorio.docx"),
    case_number: str | None = Form(None),
    include_planilha_images: bool = Form(False),
):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in content.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)

    if include_planilha_images and case_number:
        if not PDF2IMAGE_AVAILABLE:
            print("[AVISO] pdf2image indisponível. Nenhuma imagem inserida.")
        else:
            job = next(
                (j for j in JOBS.values() if j.get("case_number") == case_number),
                None,
            )
            if job:
                planilha_pages = sorted(
                    set(job.get("meta", {}).get("planilha_pages") or [])
                )
                file_path = job.get("file_path")
                if planilha_pages and file_path and os.path.exists(file_path):
                    try:
                        print(f"[INFO] Gerando imagens das páginas {planilha_pages}...")
                        images = convert_from_path(file_path)
                        doc.add_page_break()
                        doc.add_heading("Anexos – Planilhas e Bloqueios Relevantes", level=1)
                        for p in planilha_pages:
                            if 1 <= p <= len(images):
                                img_bytes = io.BytesIO()
                                images[p - 1].save(img_bytes, format="PNG")
                                img_bytes.seek(0)
                                doc.add_paragraph(f"Planilha / demonstrativo – pág. {p}")
                                doc.add_picture(img_bytes, width=Inches(6.0))
                                doc.add_paragraph("")
                    except Exception as e:
                        print(f"[AVISO] Falha ao inserir imagens no DOCX: {e}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
